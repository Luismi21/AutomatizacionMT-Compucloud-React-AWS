import json
import os
import datetime
import boto3
import base64
import mammoth
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# --- CONFIGURACIÓN ---
DOWNLOAD_BUCKET = 'memoria-tecnica-documentos-generados-123' 

# --- NUEVO: Configuración de la Plantilla en S3 ---
TEMPLATE_BUCKET = 'memoria-tecnica-documentos-generados-123' # Puede ser el mismo bucket u otro
TEMPLATE_KEY = 'plantilla/plantilla.docx' # La ruta dentro del bucket S3
LOCAL_TEMPLATE_PATH = '/tmp/plantilla.docx' # Donde se descargará en Lambda

s3_client = boto3.client('s3')

# --- LÓGICA DE LAMBDA (Reemplaza @app.route) ---
def lambda_handler(event, context):
    
    # Verificación de Bucket S3 (eliminada la condición incorrecta)
    # Puedes mantener esta verificación si quieres asegurarte que no sea el placeholder
    # if DOWNLOAD_BUCKET == '!!! REEMPLAZA-ESTO-CON-TU-BUCKET-S3-PRIVADO !!!':
    #    print("ERROR...") etc.

    try:
        # --- NUEVO: Descargar la plantilla desde S3 ---
        print(f"Descargando plantilla desde s3://{TEMPLATE_BUCKET}/{TEMPLATE_KEY} a {LOCAL_TEMPLATE_PATH}")
        try:
            s3_client.download_file(TEMPLATE_BUCKET, TEMPLATE_KEY, LOCAL_TEMPLATE_PATH)
            print("Plantilla descargada exitosamente.")
            template_path_to_use = LOCAL_TEMPLATE_PATH
        except Exception as template_error:
            print(f"ADVERTENCIA: No se pudo descargar la plantilla desde S3: {template_error}. Se intentará usar una local si existe, o crear documento en blanco.")
            # Fallback a plantilla local si existe, o None si no
            template_path_to_use = 'plantilla.docx' if os.path.exists('plantilla.docx') else None
        # --- FIN DEL NUEVO BLOQUE ---

        # 1. Obtener el archivo .json de la solicitud
        file_content = base64.b64decode(event['body'])
        input_json_path = '/tmp/estado_infraestructura.json'
        with open(input_json_path, 'wb') as f:
            f.write(file_content)

        # 2. Definir nombres de archivo de salida
        timestamp = datetime.datetime.now().strftime("%Y%m%d-%H%M%S")
        output_filename = f"Memoria_Tecnica_{timestamp}.docx"
        output_docx_path = f"/tmp/{output_filename}"

        # 3. Reutilizar tu lógica de generación
        # Ahora pasamos la ruta local donde se descargó (o None)
        generate_document_from_json(input_json_path, output_docx_path, template_path_to_use) 
        
        # 4. Generar la vista previa de HTML
        with open(output_docx_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html_preview = result.value

        # 5. Subir el .docx generado a S3
        s3_key = f"generados/{output_filename}"
        print(f"Subiendo documento generado a s3://{DOWNLOAD_BUCKET}/{s3_key}")
        s3_client.upload_file(output_docx_path, DOWNLOAD_BUCKET, s3_key)
        print("Documento subido exitosamente.")

        # 6. Generar una URL de descarga firmada (válida por 1 hora)
        download_url = s3_client.generate_presigned_url(
            'get_object',
            Params={'Bucket': DOWNLOAD_BUCKET, 'Key': s3_key},
            ExpiresIn=3600
        )

        # 7. Devolver la respuesta a React
        return {
            'statusCode': 200,
            'headers': {
                'Access-Control-Allow-Origin': '*', 
                'Access-Control-Allow-Headers': 'Content-Type',
                'Access-Control-Allow-Methods': 'POST,OPTIONS'
            },
            'body': json.dumps({
                'html_preview': html_preview,
                'download_url': download_url
            })
        }

    except Exception as e:
        print("!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
        print("!!! ERROR EN LAMBDA. TRACEBACK ABAJO: !!!")
        print("!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
        return {
            'statusCode': 500,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Access-Control-Allow-Headers': 'Content-Type',
                'Access-Control-Allow-Methods': 'POST,OPTIONS'
            },
            'body': json.dumps({'error': f"Error interno del servidor: {str(e)}"})
        }

# --- FUNCIONES DE AYUDA Y CREACIÓN DE TABLAS ---

# ... (El código de todas tus funciones de ayuda como prevent_table_split, find_resources_in_module, etc., sigue aquí SIN CAMBIOS) ...
def prevent_table_split(table):
    """Aplica propiedades a una tabla para evitar cortes extraños entre páginas."""
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        trPr.cantSplit = True
    for row in table.rows[:-1]:
        trPr = row._tr.get_or_add_trPr()
        trPr.keepNext = True

def find_resources_in_module(module, resource_type):
    """Busca recursivamente recursos de un tipo específico en un módulo y sus submódulos."""
    found_resources = []
    if 'resources' in module:
        for resource in module['resources']:
            if resource.get('type') == resource_type:
                found_resources.append(resource)
    if 'child_modules' in module:
        for child_module in module['child_modules']:
            found_resources.extend(find_resources_in_module(child_module, resource_type))
    return found_resources

def create_ec2_table(document, ec2_instance):
    values = ec2_instance.get('values', {})
    tags = values.get('tags', {})
    root_block_device = values.get('root_block_device', [{}])[0]
    heading = document.add_heading('Servidor de Cómputo (EC2)', level=1)
    heading.paragraph_format.keep_with_next = True
    table = document.add_table(rows=13, cols=6)
    table.style = 'Table Grid'
    table.cell(0, 0).merge(table.cell(6, 0)).text = tags.get('Name', 'Servidor EC2')
    table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(7, 0).merge(table.cell(9, 0)).text = 'RED'
    table.cell(7, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(7, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(10, 0).merge(table.cell(10, 5)).text = 'ALMACENAMIENTO'
    table.cell(10, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_header = table.cell(0, 1)
    cell_header.merge(table.cell(0, 5))
    cell_header.text = 'Características'
    cell_header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading_elm = parse_xml(r'<w:shd {} w:fill="00A9ED"/>'.format(nsdecls('w')))
    cell_header._tc.get_or_add_tcPr().append(shading_elm)
    fields = ["Instance ID", "Server Name", "Sistema Operativo", "Región Server", "Familia", "Key Pair Asociada", "Subred", "IP Privada", "IP Publica"]
    field_values = [
        values.get('id', 'N/A'),
        tags.get('Name', 'N/A'),
        f"Desde AMI: {values.get('ami', 'N/A')}",
        values.get('availability_zone', 'N/A').rsplit('-', 1)[0],
        values.get('instance_type', 'N/A'),
        values.get('key_name', 'N/A'),
        values.get('subnet_id', 'N/A'),
        values.get('private_ip', 'N/A'),
        values.get('public_ip', 'N/A') or 'No Asignada'
    ]
    for i, field in enumerate(fields):
        row_index = i + 1
        table.cell(row_index, 1).text = field
        value_cell = table.cell(row_index, 2)
        value_cell.merge(table.cell(row_index, 5))
        value_cell.text = field_values[i]
    storage_headers = ["ID Volumen", "Ruta", "Size (GB)", "Type", "IOPS", "Throughput"]
    for i, header in enumerate(storage_headers):
        table.cell(11, i).text = header
    storage_values = [
        root_block_device.get('volume_id', 'N/A'),
        root_block_device.get('device_name', 'N/A'),
        str(root_block_device.get('volume_size', 'N/A')),
        root_block_device.get('volume_type', 'N/A'),
        str(root_block_device.get('iops', 'N/A')),
        str(root_block_device.get('throughput', 'N/A'))
    ]
    for i, value in enumerate(storage_values):
        table.cell(12, i).text = value
    prevent_table_split(table)
    document.add_paragraph('\n')

def create_alb_table(document, alb, listeners, attachments, subnets_map):
    alb_values = alb.get('values', {})
    heading = document.add_heading('Balanceador de Carga de Aplicación (ALB)', level=1)
    heading.paragraph_format.keep_with_next = True
    esquema = "Internal" if alb_values.get('internal') else "Internet-facing"
    availability_zones = []
    for subnet_id in alb_values.get('subnets', []):
        if subnet_id in subnets_map:
            az = subnets_map.get(subnet_id, {}).get('values', {}).get('availability_zone', 'N/A')
            az_id = subnets_map.get(subnet_id, {}).get('values', {}).get('availability_zone_id', 'N/A')
            availability_zones.append(f"{az} ({az_id})")
    caracteristicas_data = {
        "Nombre": alb_values.get('name', 'N/A'),
        "Tipo": alb_values.get('load_balancer_type', 'N/A').capitalize(),
        "Esquema": esquema,
        "VPC": alb_values.get('vpc_id', 'N/A'),
        "Zonas de disponibilidad": "\n".join(availability_zones),
        "DNS name": alb_values.get('dns_name', 'N/A')
    }
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells_caract = table.rows[0].cells
    cell_header_caract = hdr_cells_caract[1]
    cell_header_caract.merge(hdr_cells_caract[3])
    cell_header_caract.text = 'Características'
    cell_header_caract.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading_caract = parse_xml(r'<w:shd {} w:fill="00A9ED"/>'.format(nsdecls('w')))
    cell_header_caract._tc.get_or_add_tcPr().append(shading_caract)
    for key, value in caracteristicas_data.items():
        row_cells = table.add_row().cells
        row_cells[1].text = key
        value_cell = row_cells[2]
        value_cell.merge(row_cells[3])
        value_cell.text = value
    listener_hdr_row = table.add_row().cells
    listener_hdr_cell = listener_hdr_row[0]
    listener_hdr_cell.merge(listener_hdr_row[3])
    listener_hdr_cell.text = 'Listeners'
    listener_hdr_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    listener_subhdr_row = table.add_row().cells
    listener_headers = ["", "Protocol:Port", "Redirect to", "Target"]
    for i, header_text in enumerate(listener_headers):
        listener_subhdr_row[i].text = header_text
    for listener in listeners:
        row_cells = table.add_row().cells
        listener_values = listener.get('values', {})
        action = listener_values.get('default_action', [{}])[0]
        row_cells[1].text = f"{listener_values.get('protocol', 'N/A')}:{listener_values.get('port', 'N/A')}"
        redirect_to, target = "N/A", "N/A"
        if action.get('type') == 'forward':
            tg_arn = action.get('forward', [{}])[0].get('target_group', [{}])[0].get('arn', '')
            if tg_arn in attachments:
                target = ", ".join(attachments[tg_arn])
                redirect_to = tg_arn.split('/')[-2]
        elif action.get('type') == 'redirect':
            redirect_to = f"Redirect ({action.get('redirect', [{}])[0].get('status_code', 'N/A')})"
            target = f"Port {action.get('redirect', [{}])[0].get('port', 'N/A')}"
        row_cells[2].text = redirect_to
        row_cells[3].text = target
    cell_title = table.cell(0, 0)
    cell_title.merge(table.cell(len(caracteristicas_data), 0))
    cell_title.text = alb_values.get('name', 'ALB')
    cell_title.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_title.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    prevent_table_split(table)
    document.add_paragraph('\n')

def create_rds_table(document, rds_instance):
    values = rds_instance.get('values', {})
    heading = document.add_heading('Base de Datos Relacional (RDS)', level=1)
    heading.paragraph_format.keep_with_next = True
    table = document.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    cell_icon = table.cell(0, 0)
    cell_icon.merge(table.cell(7, 0))
    engine = values.get('engine', 'RDS').capitalize()
    cell_icon.text = f"Amazon {engine}"
    cell_icon.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_icon.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_header = table.cell(0, 1)
    cell_header.merge(table.cell(0, 2))
    cell_header.text = 'Características'
    cell_header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading_elm = parse_xml(r'<w:shd {} w:fill="00A9ED"/>'.format(nsdecls('w')))
    cell_header._tc.get_or_add_tcPr().append(shading_elm)
    fields = {
        "DB Identifier": values.get('identifier', 'N/A'),
        "Motor": f"{values.get('engine', 'N/A')} {values.get('engine_version', '')}",
        "Tamaño": values.get('instance_class', 'N/A'),
        "Rol": "Writer Instance" if not values.get('replicate_source_db') else "Replica Instance",
        "Región Server": values.get('availability_zone', 'N/A').rsplit('-', 1)[0],
        "Endpoint": values.get('endpoint', 'N/A'),
        "Usuario master": values.get('username', 'N/A')
    }
    row_index = 1
    for key, value in fields.items():
        table.cell(row_index, 1).text = key
        table.cell(row_index, 2).text = str(value)
        row_index += 1
    prevent_table_split(table)
    document.add_paragraph('\n')

def create_vpc_table(document, vpc, route_tables_info):
    vpc_values = vpc.get('values', {})
    vpc_tags = vpc_values.get('tags', {})
    heading = document.add_heading('Red Privada Virtual (VPC)', level=1)
    heading.paragraph_format.keep_with_next = True
    table = document.add_table(rows=9, cols=3)
    table.style = 'Table Grid'
    cell_icon = table.cell(0, 0)
    cell_icon.merge(table.cell(8, 0))
    cell_icon.text = "Amazon VPC"
    cell_icon.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_icon.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_header_caract = table.cell(0, 1)
    cell_header_caract.merge(table.cell(0, 2))
    cell_header_caract.text = 'Características'
    cell_header_caract.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading_caract = parse_xml(r'<w:shd {} w:fill="00A9ED"/>'.format(nsdecls('w')))
    cell_header_caract._tc.get_or_add_tcPr().append(shading_caract)
    caracteristicas_vpc = {
        "VPC ID": vpc_values.get('id', 'N/A'),
        "Nombre vpc": vpc_tags.get('Name', 'N/A'),
        "CIDR IPv4": vpc_values.get('cidr_block', 'N/A')
    }
    row_index = 1
    for key, value in caracteristicas_vpc.items():
        table.cell(row_index, 1).text = key
        table.cell(row_index, 2).text = value
        row_index += 1
    cell_header_rt = table.cell(4, 1)
    cell_header_rt.merge(table.cell(4, 2))
    cell_header_rt.text = 'Tablas de Ruteo Asociadas'
    cell_header_rt.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt_map_local = { # Renombrado para evitar conflicto con rt_map global
        "Predeterminada": route_tables_info.get("Default", "N/A"),
        "Publica": route_tables_info.get("Public", "N/A"),
        "Privada": route_tables_info.get("Private", "N/A"),
        "RDS": route_tables_info.get("RDS", "N/A")
    }
    row_index = 5
    for key, value in rt_map_local.items():
        table.cell(row_index, 1).text = key
        table.cell(row_index, 2).text = value
        row_index += 1
    prevent_table_split(table)
    document.add_paragraph('\n')

def create_all_subnets_table(document, all_subnets, associations_map, rt_map):
    if not all_subnets: return
    heading = document.add_heading('Subredes (Subnets)', level=1)
    heading.paragraph_format.keep_with_next = True
    table = document.add_table(rows=2, cols=5)
    table.style = 'Table Grid'
    cell_header = table.cell(0, 0)
    cell_header.merge(table.cell(0, 4))
    cell_header.text = 'Características'
    cell_header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading_elm = parse_xml(r'<w:shd {} w:fill="00A9ED"/>'.format(nsdecls('w')))
    cell_header._tc.get_or_add_tcPr().append(shading_elm)
    column_headers = ["VPC ID", "Tabla de ruteo asociada", "Nombre subred", "CIDR", "AZ"]
    for i, text in enumerate(column_headers):
        table.cell(1, i).text = text
        table.cell(1, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for subnet in all_subnets:
        row_cells = table.add_row().cells
        values = subnet.get('values', {})
        tags = values.get('tags', {})
        subnet_id = values.get('id', 'N/A')
        rt_id = associations_map.get(subnet_id, None)
        route_table_name = rt_map.get(rt_id, "N/A (Principal)")
        row_cells[0].text = values.get('vpc_id', 'N/A')
        row_cells[1].text = route_table_name
        row_cells[2].text = tags.get('Name', 'N/A')
        row_cells[3].text = values.get('cidr_block', 'N/A')
        row_cells[4].text = values.get('availability_zone', 'N/A')
    cell_vpc = table.cell(2, 0)
    cell_vpc.merge(table.cell(len(all_subnets) + 1, 0))
    # Asegúrate que all_subnets no esté vacío antes de acceder al primer elemento
    if all_subnets:
        cell_vpc.text = all_subnets[0].get('values', {}).get('vpc_id', 'N/A')
    else:
         cell_vpc.text = 'N/A' # O algún valor por defecto
    cell_vpc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_vpc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    prevent_table_split(table)
    document.add_paragraph('\n')

def create_route_table_section(document, route_table, igw_map, nat_map):
    rt_values = route_table.get('values', {})
    rt_tags = rt_values.get('tags', {})
    rt_name = rt_tags.get('Name', 'N/A')
    heading = document.add_heading(f'Tabla de Ruteo: {rt_name}', level=2)
    heading.paragraph_format.keep_with_next = True
    table = document.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    cell_header_caract = table.cell(0, 1)
    cell_header_caract.merge(table.cell(0, 2))
    cell_header_caract.text = 'Características'
    cell_header_caract.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading_caract = parse_xml(r'<w:shd {} w:fill="00A9ED"/>'.format(nsdecls('w')))
    cell_header_caract._tc.get_or_add_tcPr().append(shading_caract)
    table.cell(1, 1).text = "VPC ID"
    table.cell(1, 2).text = rt_values.get('vpc_id', 'N/A')
    table.cell(2, 1).text = "Nombre Tabla"
    table.cell(2, 2).text = rt_name
    cell_header_routes = table.cell(3, 1)
    cell_header_routes.merge(table.cell(3, 2))
    cell_header_routes.text = 'Rutas'
    cell_header_routes.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(4, 1).text = "Destino"
    table.cell(4, 2).text = "Target"
    routes = rt_values.get('route', [])
    for route in routes:
        row_cells = table.add_row().cells
        row_cells[1].text = route.get('cidr_block') or route.get('ipv6_cidr_block', 'N/A')
        target = "N/A"
        if route.get('gateway_id'):
            gw_id = route.get('gateway_id')
            if gw_id == 'local':
                target = 'Local'
            elif gw_id.startswith('igw-'):
                target = f"IGW: {igw_map.get(gw_id, gw_id)}"
            else:
                target = gw_id
        elif route.get('nat_gateway_id'):
            nat_id = route.get('nat_gateway_id')
            target = f"NAT GW: {nat_map.get(nat_id, nat_id)}"
        row_cells[2].text = target
    cell_icon = table.cell(0, 0)
    cell_icon.merge(table.cell(len(table.rows) - 1, 0))
    cell_icon.text = "Rutas"
    cell_icon.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_icon.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    prevent_table_split(table)
    document.add_paragraph('\n')

def create_igw_section(document, igw):
    igw_values = igw.get('values', {})
    igw_tags = igw_values.get('tags', {})
    igw_name = igw_tags.get('Name', 'N/A')
    heading = document.add_heading(f"Internet Gateway: {igw_name}", level=2)
    heading.paragraph_format.keep_with_next = True
    table = document.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    cell_icon = table.cell(0, 0)
    cell_icon.merge(table.cell(3, 0))
    cell_icon.text = "IGW"
    cell_icon.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_icon.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_header_caract = table.cell(0, 1)
    cell_header_caract.merge(table.cell(0, 2))
    cell_header_caract.text = 'Características'
    cell_header_caract.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading_caract = parse_xml(r'<w:shd {} w:fill="00A9ED"/>'.format(nsdecls('w')))
    cell_header_caract._tc.get_or_add_tcPr().append(shading_caract)
    table.cell(1, 1).text = "VPC ID"
    table.cell(1, 2).text = igw_values.get('vpc_id', 'N/A')
    table.cell(2, 1).text = "Nombre IGW"
    table.cell(2, 2).text = igw_name
    table.cell(3, 1).text = "IGW ID"
    table.cell(3, 2).text = igw_values.get('id', 'N/A')
    prevent_table_split(table)
    document.add_paragraph('\n')

def create_nat_gateway_table(document, nat_gateway, subnets_map):
    nat_values = nat_gateway.get('values', {})
    nat_tags = nat_values.get('tags', {})
    heading = document.add_heading(f"NAT Gateway: {nat_tags.get('Name', 'N/A')}", level=2)
    heading.paragraph_format.keep_with_next = True
    table = document.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    icon_cell = table.cell(0, 0)
    icon_cell.merge(table.cell(4, 0))
    icon_cell.text = "NAT Gateway"
    icon_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    icon_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    header_caract_cell = table.cell(0, 1)
    header_caract_cell.merge(table.cell(0, 2))
    header_caract_cell.text = 'Características'
    header_caract_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading_caract = parse_xml(r'<w:shd {} w:fill="00A9ED"/>'.format(nsdecls('w')))
    header_caract_cell._tc.get_or_add_tcPr().append(shading_caract)
    subnet_id = nat_values.get('subnet_id', 'N/A')
    subnet_info = subnets_map.get(subnet_id, {}).get('values', {})
    subnet_name = subnet_info.get('tags', {}).get('Name', subnet_id)
    subnet_az = subnet_info.get('availability_zone', 'N/A').rsplit('-', 1)[-1] if subnet_info.get('availability_zone') else 'N/A'
    data_rows = [
        ("VPC ID", subnets_map.get(subnet_id, {}).get('values', {}).get('vpc_id', 'N/A')),
        ("Subnet", f"{subnet_id} / {subnet_name} - AZ {subnet_az}"),
        ("Nombre NATGW", nat_tags.get('Name', 'N/A')),
        ("NATGW ID", nat_values.get('id', 'N/A'))
    ]
    for i, (label, value) in enumerate(data_rows, start=1):
        table.cell(i, 1).text = label
        table.cell(i, 2).text = value
    prevent_table_split(table)
    document.add_paragraph('\n')

def create_target_group_table(document, target_group, attachments):
    tg_values = target_group.get('values', {})
    tg_name = tg_values.get('name', 'N/A')
    heading = document.add_heading(f'Grupo de Destino: {tg_name}', level=2)
    heading.paragraph_format.keep_with_next = True
    table = document.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    header_caract_cell = table.cell(0, 0)
    header_caract_cell.merge(table.cell(0, 1))
    header_caract_cell.text = 'Características'
    header_caract_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading_caract = parse_xml(r'<w:shd {} w:fill="00A9ED"/>'.format(nsdecls('w')))
    header_caract_cell._tc.get_or_add_tcPr().append(shading_caract)
    data_rows = [
        ("Nombre", tg_name),
        ("Tipo de destino", tg_values.get('target_type', 'N/A').capitalize()),
        ("Protocolo", tg_values.get('protocol', 'N/A')),
        ("Puerto", str(tg_values.get('port', 'N/A')))
    ]
    for i, (label, value) in enumerate(data_rows, start=1):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
    header_instances_cell = table.cell(5, 0)
    header_instances_cell.merge(table.cell(5, 1))
    header_instances_cell.text = 'Instancias Asociadas'
    header_instances_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tg_arn = tg_values.get('arn')
    associated_instance_ids = attachments.get(tg_arn, [])
    if associated_instance_ids:
        for instance_id in associated_instance_ids:
            row_cells = table.add_row().cells
            row_cells[0].merge(row_cells[1])
            row_cells[0].text = instance_id
    else:
        row_cells = table.add_row().cells
        row_cells[0].merge(row_cells[1])
        row_cells[0].text = "No hay instancias asociadas"
    prevent_table_split(table)
    document.add_paragraph('\n')

def create_kms_table(document, kms_key, aliases_map):
    kms_values = kms_key.get('values', {})
    key_id = kms_values.get('id')
    alias = aliases_map.get(key_id, 'N/A')
    heading = document.add_heading('Key Management Services (KMS)', level=2)
    heading.paragraph_format.keep_with_next = True
    document.add_paragraph(alias)
    table = document.add_table(rows=3, cols=3) # Changed rows to 3
    table.style = 'Table Grid'
    side_title_cell = table.cell(0, 0)
    side_title_cell.merge(table.cell(2, 0)) # Changed merge to 2
    side_title_cell.text = 'Claves administradas'
    side_title_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    side_title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    header_cells = table.rows[0].cells
    header_cells[1].text = 'Alias'
    header_cells[2].text = alias
    data_rows = [
        ("ID de la Clave", key_id),
        ("Descripción", kms_values.get('description', 'N/A')),
    ]
    table.cell(1, 1).text = data_rows[0][0]
    table.cell(1, 2).text = data_rows[0][1]
    table.cell(2, 1).text = data_rows[1][0]
    table.cell(2, 2).text = data_rows[1][1]
    prevent_table_split(table)
    document.add_paragraph('\n')


# --- LÓGICA PRINCIPAL (Llamada por el handler) ---

def generate_document_from_json(input_json_path, output_docx_path, template_path):
    """Función que orquesta la creación del documento de Word."""
    
    # Corrección: El JSON del usuario es UTF-16
    with open(input_json_path, 'r', encoding='utf-16') as f:
        data = json.load(f)

    root_module = data.get('values', {}).get('root_module', {})
    
    # --- CAMBIO IMPORTANTE: Manejo de Plantilla Opcional ---
    # CAMBIO: Usar template_path que viene como argumento (puede ser None)
    try:
        if template_path and os.path.exists(template_path): # Verifica si existe y no es None
            document = Document(template_path)
            print("Usando plantilla descargada/local encontrada.")
        else:
            print("Advertencia: No se encontró plantilla válida. Creando documento en blanco.")
            document = Document()
    except Exception as e:
        print(f"Error al cargar la plantilla desde '{template_path}': {e}. Creando documento en blanco.")
        document = Document()
    # --- FIN DEL CAMBIO ---
    
    document.add_heading('Memoria Técnica de Infraestructura AWS', 1)
    document.add_paragraph('Este documento contiene un resumen detallado...')
    document.add_paragraph('')

    all_subnets = find_resources_in_module(root_module, 'aws_subnet')
    all_route_tables = find_resources_in_module(root_module, 'aws_route_table')
    all_associations = find_resources_in_module(root_module, 'aws_route_table_association')
    all_igws = find_resources_in_module(root_module, 'aws_internet_gateway')
    all_nat_gws = find_resources_in_module(root_module, 'aws_nat_gateway')
    all_tgs = find_resources_in_module(root_module, 'aws_lb_target_group')
    all_kms_keys = find_resources_in_module(root_module, 'aws_kms_key')
    all_kms_aliases = find_resources_in_module(root_module, 'aws_kms_alias')

    # Corrección para evitar error si alguna lista está vacía
    subnet_map = {s['values']['id']: s for s in all_subnets if 'values' in s and 'id' in s['values']}
    rt_map = {rt['values']['id']: rt['values'].get('tags', {}).get('Name', rt['values']['id']) for rt in all_route_tables if 'values' in rt and 'id' in rt['values']}
    associations_map = {assoc['values']['subnet_id']: assoc['values']['route_table_id'] for assoc in all_associations if 'values' in assoc and 'subnet_id' in assoc['values']}
    igw_map = {igw['values']['id']: igw['values'].get('tags', {}).get('Name', igw['values']['id']) for igw in all_igws if 'values' in igw and 'id' in igw['values']}
    nat_map = {nat['values']['id']: nat['values'].get('tags', {}).get('Name', nat['values']['id']) for nat in all_nat_gws if 'values' in nat and 'id' in nat['values']}
    aliases_map = {alias['values']['target_key_id']: alias['values'].get('name', '').replace('alias/', '') for alias in all_kms_aliases if 'values' in alias and 'target_key_id' in alias['values']}


    all_vpcs = find_resources_in_module(root_module, 'aws_vpc')
    if all_vpcs:
        for vpc in all_vpcs:
            # Asegúrate que 'values' y 'id' existen antes de usarlos
            if 'values' not in vpc or 'id' not in vpc['values']:
                print(f"Advertencia: VPC encontrada sin 'values' o 'id'. Saltando: {vpc}")
                continue
            route_tables_info = {}
            vpc_id = vpc['values']['id']
            default_rt_id = vpc['values'].get('main_route_table_id')
            if default_rt_id in rt_map:
                route_tables_info["Default"] = rt_map[default_rt_id]
            
            # Filtro más seguro para subnet_name_map
            subnet_name_map = {s['values']['id']: s['values'].get('tags', {}).get('Name', '').lower() 
                               for s in all_subnets 
                               if 'values' in s and 'id' in s['values'] and s['values'].get('vpc_id') == vpc_id}

            for assoc in all_associations:
                if 'values' not in assoc: continue # Asegurarse que 'values' existe
                subnet_id = assoc['values'].get('subnet_id')
                if subnet_id in subnet_name_map:
                    subnet_name = subnet_name_map[subnet_id]
                    rt_id = assoc['values'].get('route_table_id')
                    rt_name = rt_map.get(rt_id, "N/A")
                    
                    if "public" in subnet_name and "Public" not in route_tables_info:
                        route_tables_info["Public"] = rt_name
                    elif "private" in subnet_name and "Private" not in route_tables_info:
                        route_tables_info["Private"] = rt_name
                    elif "rds" in subnet_name and "RDS" not in route_tables_info:
                        route_tables_info["RDS"] = rt_name
            
            create_vpc_table(document, vpc, route_tables_info)
    else: 
        print("ℹ️ No se encontraron VPCs.")

    if all_subnets:
        all_subnets.sort(key=lambda s: s.get('values', {}).get('tags', {}).get('Name', ''))
        create_all_subnets_table(document, all_subnets, associations_map, rt_map)
    else:
        print("ℹ️ No se encontraron Subredes.")

    if all_route_tables:
        document.add_heading('Sección de Ruteo', level=1)
        all_route_tables.sort(key=lambda s: s.get('values', {}).get('tags', {}).get('Name', ''))
        for rt in all_route_tables:
             # Asegúrate que 'values' existe antes de llamar a la función
            if 'values' in rt:
                create_route_table_section(document, rt, igw_map, nat_map)
            else:
                 print(f"Advertencia: Tabla de ruteo encontrada sin 'values'. Saltando: {rt}")
    else:
        print("ℹ️ No se encontraron Tablas de Ruteo.")

    if all_igws:
        document.add_heading('Gateways de Internet', level=1)
        for igw in all_igws:
            if 'values' in igw:
                create_igw_section(document, igw)
            else:
                 print(f"Advertencia: IGW encontrado sin 'values'. Saltando: {igw}")
    else:
        print("ℹ️ No se encontraron Gateways de Internet.")

    if all_nat_gws:
        document.add_heading('Gateways NAT', level=1)
        all_nat_gws.sort(key=lambda s: s.get('values', {}).get('tags', {}).get('Name', ''))
        for nat_gw in all_nat_gws:
             if 'values' in nat_gw:
                create_nat_gateway_table(document, nat_gw, subnet_map)
             else:
                  print(f"Advertencia: NAT GW encontrado sin 'values'. Saltando: {nat_gw}")
    else:
        print("ℹ️ No se encontraron NAT Gateways.")

    ec2_instances = find_resources_in_module(root_module, 'aws_instance')
    if ec2_instances:
        for instance in ec2_instances: 
            if 'values' in instance:
                create_ec2_table(document, instance)
            else:
                 print(f"Advertencia: Instancia EC2 encontrada sin 'values'. Saltando: {instance}")
    else: print("ℹ️ No se encontraron instancias EC2.")

    all_albs = find_resources_in_module(root_module, 'aws_lb')
    if all_albs:
        all_listeners = find_resources_in_module(root_module, 'aws_lb_listener')
        all_tg_attachments = find_resources_in_module(root_module, 'aws_lb_target_group_attachment')
        tg_attachments_map = {}
        for att in all_tg_attachments:
             # Asegúrate que 'values' y las claves necesarias existen
            if 'values' in att and 'target_group_arn' in att['values'] and 'target_id' in att['values']:
                tg_arn = att['values']['target_group_arn']
                if tg_arn not in tg_attachments_map:
                    tg_attachments_map[tg_arn] = []
                tg_attachments_map[tg_arn].append(att['values']['target_id'])
            else:
                 print(f"Advertencia: Adjunto de TG encontrado con estructura inesperada. Saltando: {att}")


        for alb in all_albs:
             if 'values' in alb and 'arn' in alb['values']:
                alb_arn = alb['values']['arn']
                # Filtro más seguro para listeners
                relevant_listeners = [l for l in all_listeners 
                                      if 'values' in l and l['values'].get('load_balancer_arn') == alb_arn]
                create_alb_table(document, alb, relevant_listeners, tg_attachments_map, subnet_map)
             else:
                  print(f"Advertencia: ALB encontrado sin 'values' o 'arn'. Saltando: {alb}")

    else: print("ℹ️ No se encontraron Balanceadores de Carga.")

    if all_tgs:
        document.add_heading('Grupos de Destino (Target Groups)', level=1)
        all_tgs.sort(key=lambda s: s.get('values', {}).get('name', ''))
        # Crear tg_attachments_map aquí si no se creó antes y si es necesario
        if 'tg_attachments_map' not in locals() and any('values' in tg for tg in all_tgs): # Solo si hay TGs con 'values'
            all_tg_attachments = find_resources_in_module(root_module, 'aws_lb_target_group_attachment')
            tg_attachments_map = {}
            for att in all_tg_attachments:
                 if 'values' in att and 'target_group_arn' in att['values'] and 'target_id' in att['values']:
                    tg_arn = att['values']['target_group_arn']
                    if tg_arn not in tg_attachments_map:
                        tg_attachments_map[tg_arn] = []
                    tg_attachments_map[tg_arn].append(att['values']['target_id'])
                 else:
                     print(f"Advertencia: Adjunto de TG encontrado con estructura inesperada. Saltando: {att}")
            
        for tg in all_tgs:
             if 'values' in tg:
                # Asegúrate que tg_attachments_map exista antes de pasarla
                current_attachments = tg_attachments_map if 'tg_attachments_map' in locals() else {}
                create_target_group_table(document, tg, current_attachments)
             else:
                 print(f"Advertencia: Target Group encontrado sin 'values'. Saltando: {tg}")
    else:
        print("ℹ️ No se encontraron Target Groups.")

    rds_instances = find_resources_in_module(root_module, 'aws_db_instance')
    if rds_instances:
        for instance in rds_instances: 
            if 'values' in instance:
                create_rds_table(document, instance)
            else:
                 print(f"Advertencia: Instancia RDS encontrada sin 'values'. Saltando: {instance}")

    else: print("ℹ️ No se encontraron instancias RDS.")

    if all_kms_keys:
        document.add_heading('Servicios de Gestión de Claves (KMS)', level=1)
        for kms_key in all_kms_keys:
             if 'values' in kms_key:
                create_kms_table(document, kms_key, aliases_map)
             else:
                  print(f"Advertencia: Clave KMS encontrada sin 'values'. Saltando: {kms_key}")
    else:
        print("ℹ️ No se encontraron Claves KMS.")

    document.save(output_docx_path)

